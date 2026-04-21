"""
demo.py - AETHER Agent Demo
Demonstrates the complete DAGR pipeline with AEC verification.
"""

import argparse
import json
import os
import sys
import textwrap
import time
import threading
from datetime import datetime
from pathlib import Path

# ═══════════════════════════════════════════════════════════════════════════════
# ANSI COLORS AND CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════════

BLUE = '\033[94m'
GREEN = '\033[92m'
RED = '\033[91m'
YELLOW = '\033[93m'
CYAN = '\033[96m'
MAGENTA = '\033[95m'
WHITE = '\033[97m'
BOLD = '\033[1m'
DIM = '\033[2m'
RESET = '\033[0m'

WIDTH = 66  # Box width

# Demo query
DEMO_QUERY = """Create a professional project status report structure with executive summary, milestones table, and risk register. Include guidance on heading styles and table formatting."""

# Domain entities to extract (extends basic distill)
DOMAIN_ENTITIES = [
    "executive summary", "milestones", "risk register",
    "heading styles", "table formatting", "project status",
    "document structure", "report"
]

# Stub response for --stub mode
STUB_RESPONSE = """To create a professional project status report in Word:

1. Use Heading 1 style for "Executive Summary" at the top of your document. This section should provide a high-level overview in 2-3 paragraphs.

2. For the Milestones Table, create a 3-column table with columns for Milestone, Target Date, and Status. Always use DXA units for table widths to ensure compatibility - set both columnWidths on the table AND width on each cell for reliable rendering.

3. Apply Arial as your default font since it's universally supported across all systems and renders consistently.

4. The Risk Register should use Heading 1 style and contain a table with Risk Description, Impact, Probability, and Mitigation columns.

5. For all tables, use ShadingType.CLEAR instead of SOLID to prevent black backgrounds from appearing in some viewers.

6. Always use DXA units for table widths - WidthType.PERCENTAGE breaks rendering in Google Docs and some Word versions.

7. Keep titles black for maximum readability. Avoid using unicode bullet characters - instead use numbering config with LevelFormat.BULLET for proper list formatting.

8. After creating the file, validate it using the validation script to ensure OOXML compliance. This will give your report a professional appearance."""


# ═══════════════════════════════════════════════════════════════════════════════
# SPINNER AND PROGRESS
# ═══════════════════════════════════════════════════════════════════════════════

class Spinner:
    """Animated spinner for progress indication."""
    FRAMES = ['⠋', '⠙', '⠹', '⠸', '⠼', '⠴', '⠦', '⠧', '⠇', '⠏']

    def __init__(self, message: str):
        self.message = message
        self.running = False
        self.thread = None
        self.frame_idx = 0

    def _spin(self):
        while self.running:
            frame = self.FRAMES[self.frame_idx % len(self.FRAMES)]
            sys.stdout.write(f'\r{CYAN}{frame}{RESET} {self.message}')
            sys.stdout.flush()
            self.frame_idx += 1
            time.sleep(0.08)

    def start(self):
        self.running = True
        self.thread = threading.Thread(target=self._spin, daemon=True)
        self.thread.start()

    def stop(self, success_msg: str = None, time_ms: float = None):
        self.running = False
        if self.thread:
            self.thread.join(timeout=0.2)
        # Clear line and print result
        sys.stdout.write('\r' + ' ' * 80 + '\r')
        if success_msg:
            time_str = f"  ({time_ms:.0f}ms)" if time_ms else ""
            if time_ms and time_ms > 1000:
                time_str = f"  ({time_ms/1000:.1f}s)"
            print(f"{GREEN}✓{RESET} {success_msg}{DIM}{time_str}{RESET}")


# ═══════════════════════════════════════════════════════════════════════════════
# DISPLAY HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def print_banner():
    """Print dramatic ASCII art banner."""
    print()
    print(f"{CYAN}╔{'═'*WIDTH}╗{RESET}")
    print(f"{CYAN}║{' '*WIDTH}║{RESET}")
    print(f"{CYAN}║{BOLD}   █████╗ ███████╗████████╗██╗  ██╗███████╗██████╗       {RESET}{CYAN}║{RESET}")
    print(f"{CYAN}║{BOLD}  ██╔══██╗██╔════╝╚══██╔══╝██║  ██║██╔════╝██╔══██╗      {RESET}{CYAN}║{RESET}")
    print(f"{CYAN}║{BOLD}  ███████║█████╗     ██║   ███████║█████╗  ██████╔╝      {RESET}{CYAN}║{RESET}")
    print(f"{CYAN}║{BOLD}  ██╔══██║██╔══╝     ██║   ██╔══██║██╔══╝  ██╔══██╗      {RESET}{CYAN}║{RESET}")
    print(f"{CYAN}║{BOLD}  ██║  ██║███████╗   ██║   ██║  ██║███████╗██║  ██║      {RESET}{CYAN}║{RESET}")
    print(f"{CYAN}║{BOLD}  ╚═╝  ╚═╝╚══════╝   ╚═╝   ╚═╝  ╚═╝╚══════╝╚═╝  ╚═╝      {RESET}{CYAN}║{RESET}")
    print(f"{CYAN}║{' '*WIDTH}║{RESET}")
    print(f"{CYAN}║  {WHITE}Agent Education Calibration{RESET}  {DIM}·{RESET}  {WHITE}Skill Agent Demo{RESET}{' '*12}{CYAN}║{RESET}")
    print(f"{CYAN}║  {DIM}\"The model did not change. The skill did.\"{RESET}{' '*21}{CYAN}║{RESET}")
    print(f"{CYAN}╚{'═'*WIDTH}╝{RESET}")
    print()


def print_beat(num: int, title: str):
    """Print dramatic beat header."""
    print()
    print(f"{CYAN}{'━'*WIDTH}{RESET}")
    print(f"{CYAN}◈  BEAT {num}  ·  {title}{RESET}")
    print(f"{CYAN}{'━'*WIDTH}{RESET}")
    print()


def print_box_top():
    print(f"┌{'─'*(WIDTH-2)}┐")

def print_box_line(text: str, prefix: str = ""):
    """Print a line inside a box with proper padding."""
    content = f"{prefix}{text}"
    # Handle ANSI codes in length calculation
    visible_len = len(content.replace(BLUE, '').replace(GREEN, '').replace(RED, '')
                       .replace(YELLOW, '').replace(CYAN, '').replace(MAGENTA, '')
                       .replace(WHITE, '').replace(BOLD, '').replace(DIM, '').replace(RESET, ''))
    padding = WIDTH - 4 - visible_len
    if padding < 0:
        padding = 0
    print(f"│  {content}{' '*padding}│")

def print_box_empty():
    print(f"│{' '*(WIDTH-2)}│")

def print_box_bottom():
    print(f"└{'─'*(WIDTH-2)}┘")

def print_box_divider():
    print(f"│  {'─'*(WIDTH-6)}  │")


def render_gauge(score: float, threshold: float = 0.80, width: int = 20) -> str:
    """Render a visual score gauge."""
    filled = int(score * width)
    bar = '█' * filled + '░' * (width - filled)
    color = GREEN if score >= threshold else RED
    return f"{color}[{bar}]{RESET} {score:.3f}   threshold {threshold:.2f}"


def count_node_types(kg_nodes: list) -> dict:
    """Count nodes by type."""
    counts = {}
    for node in kg_nodes:
        ntype = node.get('@type', '')
        if 'Rule' in ntype:
            counts['Rule'] = counts.get('Rule', 0) + 1
        elif 'Technique' in ntype:
            counts['Technique'] = counts.get('Technique', 0) + 1
        elif 'AntiPattern' in ntype:
            counts['AntiPattern'] = counts.get('AntiPattern', 0) + 1
        elif 'Concept' in ntype:
            counts['Concept'] = counts.get('Concept', 0) + 1
        elif 'Tool' in ntype:
            counts['Tool'] = counts.get('Tool', 0) + 1
        elif 'Trait' in ntype:
            counts['Trait'] = counts.get('Trait', 0) + 1
    return counts


def extract_domain_entities(query: str) -> list:
    """Extract domain-specific entities from query."""
    query_lower = query.lower()
    found = []
    for entity in DOMAIN_ENTITIES:
        if entity in query_lower:
            # Title case for display
            found.append(entity.title())
    return found


def resolve_capsule_path(capsule_arg: str, examples_dir: Path) -> Path:
    """Resolve capsule path from short name or full path."""
    if Path(capsule_arg).is_dir():
        return Path(capsule_arg)
    for d in examples_dir.iterdir():
        if d.is_dir() and d.name.startswith(capsule_arg):
            return d
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX FILE CREATION (Node.js + docx npm per SKILL.md)
# ═══════════════════════════════════════════════════════════════════════════════

def create_docx_file(response: str, output_path: Path) -> tuple:
    """
    Create a .docx file using Node.js + docx npm library.
    This follows the SKILL.md specification exactly:
    - US Letter page size (12240 x 15840 DXA)
    - 1 inch margins (1440 DXA)
    - Arial font
    - DXA units for tables (not percentages)
    - ShadingType.CLEAR (not SOLID)
    - Proper heading style overrides

    Returns (success, size_kb).
    """
    import subprocess

    script_path = Path(__file__).parent / "scripts" / "generate_docx_demo.js"

    if not script_path.exists():
        # Fallback to python-docx if script missing
        return _create_docx_fallback(response, output_path)

    try:
        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Run Node.js generator
        result = subprocess.run(
            ["node", str(script_path), str(output_path)],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent,
            timeout=30,
        )

        if result.returncode != 0:
            # Try fallback
            return _create_docx_fallback(response, output_path)

        # Parse result
        import json
        try:
            output = json.loads(result.stdout.strip())
            if output.get('success'):
                return True, output.get('size_kb', 0)
        except json.JSONDecodeError:
            pass

        # Check if file was created
        if output_path.exists():
            size_kb = output_path.stat().st_size / 1024
            return True, size_kb

        return _create_docx_fallback(response, output_path)

    except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
        return _create_docx_fallback(response, output_path)


def _create_docx_fallback(response: str, output_path: Path) -> tuple:
    """Fallback to python-docx if Node.js generation fails."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return False, 0

    doc = Document()

    # Title
    title = doc.add_heading('Project Status Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report provides a comprehensive overview of the current project status, '
        'including key milestones, deliverables, and risk assessment.'
    )

    # Milestones Table
    doc.add_heading('Milestones', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Milestone'
    hdr_cells[1].text = 'Target Date'
    hdr_cells[2].text = 'Status'

    milestones = [
        ('Phase 1 Complete', '2026-Q1', 'Complete'),
        ('Phase 2 Kickoff', '2026-Q2', 'In Progress'),
        ('Final Delivery', '2026-Q4', 'Planned'),
    ]
    for i, (milestone, date, status) in enumerate(milestones, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = milestone
        row_cells[1].text = date
        row_cells[2].text = status

    # Risk Register
    doc.add_heading('Risk Register', level=1)
    risk_table = doc.add_table(rows=3, cols=4)
    risk_table.style = 'Table Grid'

    risk_hdr = risk_table.rows[0].cells
    risk_hdr[0].text = 'Risk'
    risk_hdr[1].text = 'Impact'
    risk_hdr[2].text = 'Probability'
    risk_hdr[3].text = 'Mitigation'

    risks = [
        ('Resource constraints', 'High', 'Medium', 'Cross-train team'),
        ('Scope creep', 'Medium', 'High', 'Change control process'),
    ]
    for i, (risk, impact, prob, mitigation) in enumerate(risks, 1):
        row = risk_table.rows[i].cells
        row[0].text = risk
        row[1].text = impact
        row[2].text = prob
        row[3].text = mitigation

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('Generated by AETHER docx agent v1.0.0 (fallback)').italic = True

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    size_kb = output_path.stat().st_size / 1024
    return True, size_kb


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN DEMO
# ═══════════════════════════════════════════════════════════════════════════════

def run_demo(capsule_path: Path, stub_mode: bool, no_docx: bool):
    """Run the full demo."""
    from aether import Capsule
    from aec import verify, split_statements
    from aec_concept import compile_kg, has_typed_nodes
    from llm import make_llm_fn
    from kg import load_kg, get_nodes

    # Track timing
    timing = {}
    pipeline_start = time.time()

    # Check API key if not stub mode
    if not stub_mode and not os.environ.get('ANTHROPIC_API_KEY'):
        print(f"{RED}No ANTHROPIC_API_KEY found.{RESET}")
        print(f"Run with --stub for demo mode:")
        print(f"  python demo.py --stub")
        sys.exit(1)

    print_banner()

    if stub_mode:
        print(f"{DIM}  [stub mode - set ANTHROPIC_API_KEY for live generation]{RESET}")
        print()

    # ═══════════════════════════════════════════════════════════════════════════
    # BEAT 1: THE AGENT
    # ═══════════════════════════════════════════════════════════════════════════
    print_beat(1, "THE AGENT")

    spinner = Spinner("Loading capsule...")
    spinner.start()
    load_start = time.time()

    llm_fn = make_llm_fn(provider="stub") if stub_mode else make_llm_fn(provider="anthropic")
    capsule = Capsule(str(capsule_path), llm_fn=llm_fn)

    files = capsule.files
    kg_nodes = files['kg'].get('@graph', [])
    node_counts = count_node_types(kg_nodes)
    kb_text = files['kb']
    kb_paragraphs = len([p for p in kb_text.split('\n\n') if p.strip()])

    timing['load'] = (time.time() - load_start) * 1000
    spinner.stop("Capsule loaded", timing['load'])
    print()

    # Agent Card
    print_box_top()
    print_box_line(f"{BOLD}AGENT LOADED{RESET}")
    print_box_empty()
    print_box_line(f"{BLUE}Name{RESET}     {capsule.id}")
    print_box_line(f"{BLUE}Origin{RESET}   Anthropic SKILL.md → DAG → compiled capsule")
    print_box_line(f"{BLUE}KB{RESET}       {len(kb_text):,} chars  ·  {kb_paragraphs} paragraphs of knowledge")
    print_box_line(f"{BLUE}KG{RESET}       {len(kg_nodes)} nodes compiled → {len(kg_nodes)} detectors")

    # Type breakdown
    type_line = f"         Rule({node_counts.get('Rule', 0)}) Technique({node_counts.get('Technique', 0)}) AntiPattern({node_counts.get('AntiPattern', 0)})"
    print_box_line(type_line)
    type_line2 = f"         Concept({node_counts.get('Concept', 0)}) Tool({node_counts.get('Tool', 0)}) Trait({node_counts.get('Trait', 0)})"
    print_box_line(type_line2)
    print_box_empty()
    print_box_line(f"{DIM}\"This agent started as a plain markdown file.\"{RESET}")
    print_box_bottom()

    # ═══════════════════════════════════════════════════════════════════════════
    # BEAT 2: DISTILL
    # ═══════════════════════════════════════════════════════════════════════════
    print_beat(2, "DISTILL")

    spinner = Spinner("Distilling query...")
    spinner.start()
    distill_start = time.time()

    ctx = {"input": DEMO_QUERY}
    ctx = capsule.distill(ctx)
    distill_result = ctx.get('distilled', {})

    timing['distill'] = (time.time() - distill_start) * 1000
    spinner.stop("Distill complete", timing['distill'])
    print()

    # Extract domain entities
    domain_entities = extract_domain_entities(DEMO_QUERY)
    all_entities = list(set(distill_result.get('entities', []) + domain_entities))

    print(f"  {BLUE}Intent{RESET}     {distill_result.get('intent', 'unknown')}")
    entities_str = ", ".join(all_entities) if all_entities else "none"
    print(f"  {BLUE}Entities{RESET}   {entities_str}")
    format_hint = distill_result.get('format', 'general') or 'general'
    print(f"  {BLUE}Format{RESET}     {format_hint} detected")

    # ═══════════════════════════════════════════════════════════════════════════
    # BEAT 3: AUGMENT
    # ═══════════════════════════════════════════════════════════════════════════
    print_beat(3, "AUGMENT — Knowledge Retrieved")

    spinner = Spinner("Augmenting context...")
    spinner.start()
    augment_start = time.time()

    ctx = capsule.augment(ctx)
    augment_result = ctx.get('augmented', {})

    timing['augment'] = (time.time() - augment_start) * 1000
    spinner.stop("Augment complete", timing['augment'])
    print()

    # KB Display
    kb_context = augment_result.get('kb_context', '') or augment_result.get('kb', '')
    kb_matches = augment_result.get('kb_matches', [])

    # Handle kb_context being a list or string
    if isinstance(kb_context, list):
        passages = kb_context[:3]
    elif isinstance(kb_context, str) and kb_context:
        passages = [p.strip() for p in kb_context.split('\n\n') if p.strip()][:3]
    else:
        passages = []

    if passages:
        print(f"  {BLUE}KB{RESET}  {len(passages)} passages matched from {len(kb_text):,} chars of knowledge")
        print()
        for i, passage in enumerate(passages, 1):
            # Handle passage being dict or string
            if isinstance(passage, dict):
                text = passage.get('text', passage.get('content', str(passage)))
            else:
                text = str(passage)
            # Truncate to ~120 chars and wrap
            truncated = text[:120] + "..." if len(text) > 120 else text
            wrapped = textwrap.fill(truncated, width=55, initial_indent="      ", subsequent_indent="           ")
            print(f"      {DIM}[{i}]{RESET} {wrapped[6:]}")  # Remove initial indent from first line
            print()
    else:
        print(f"  {BLUE}KB{RESET}  Searching knowledge base...")
        print()

    # KG Display
    kg_matches = augment_result.get('kg_matches', []) or augment_result.get('kg', [])
    if kg_matches:
        print(f"  {BLUE}KG{RESET}  {len(kg_matches)} nodes retrieved as context")
        print()
        for match in kg_matches[:5]:
            node_id = match.get('@id', match.get('node_id', 'unknown'))
            label = match.get('rdfs:label', match.get('label', ''))
            # Truncate label
            if len(label) > 35:
                label = label[:35] + "..."
            print(f"      {CYAN}{node_id:<28}{RESET} → {label}")
    else:
        # Show some relevant nodes from KG based on keywords
        keywords = distill_result.get('keywords', [])
        relevant = []
        for node in kg_nodes[:10]:
            label = node.get('rdfs:label', '').lower()
            if any(kw in label for kw in ['heading', 'table', 'document', 'style', 'format']):
                relevant.append(node)
        if relevant:
            print(f"  {BLUE}KG{RESET}  {len(relevant)} relevant nodes available")
            print()
            for node in relevant[:5]:
                node_id = node.get('@id', 'unknown')
                label = node.get('rdfs:label', '')
                if len(label) > 35:
                    label = label[:35] + "..."
                print(f"      {CYAN}{node_id:<28}{RESET} → {label}")

    # ═══════════════════════════════════════════════════════════════════════════
    # BEAT 4: RESPONSE
    # ═══════════════════════════════════════════════════════════════════════════
    print_beat(4, "RESPONSE")

    if stub_mode:
        spinner = Spinner("Generating response...")
        spinner.start()
        time.sleep(0.3)  # Brief pause for effect in stub mode
        response = STUB_RESPONSE
        timing['generate'] = 0
        spinner.stop("Response generated (stub)", 0)
    else:
        spinner = Spinner("Generating response...")
        spinner.start()
        generate_start = time.time()

        ctx = capsule.run(DEMO_QUERY)
        response = ctx.get('generated', '')

        timing['generate'] = (time.time() - generate_start) * 1000
        spinner.stop("Response generated", timing['generate'])

    print()

    # Response box with word wrapping
    print_box_top()
    print_box_line(f"{BOLD}AGENT RESPONSE{RESET}")
    print_box_divider()
    print_box_empty()

    # Word wrap the response
    wrapper = textwrap.TextWrapper(width=WIDTH-6, break_long_words=False, break_on_hyphens=False)
    for paragraph in response.split('\n'):
        if paragraph.strip():
            wrapped_lines = wrapper.wrap(paragraph)
            for line in wrapped_lines:
                print_box_line(line)
        else:
            print_box_empty()

    print_box_empty()
    print_box_bottom()

    # ═══════════════════════════════════════════════════════════════════════════
    # BEAT 5: AEC VERIFICATION
    # ═══════════════════════════════════════════════════════════════════════════
    print_beat(5, "AEC VERIFICATION")

    # Compile KG
    spinner = Spinner("Compiling KG...")
    spinner.start()
    compile_start = time.time()

    compiled_kg = compile_kg(kg_nodes) if has_typed_nodes(kg_nodes) else None

    compile_ms = (time.time() - compile_start) * 1000
    spinner.stop(f"Compiled {len(kg_nodes)} nodes → {len(compiled_kg['detectors']) if compiled_kg else 0} detectors", compile_ms)
    print()

    # Run AEC
    spinner = Spinner("Verifying statements...")
    spinner.start()
    aec_start = time.time()

    statements = split_statements(response)
    aec_result = verify(response, kg_nodes, threshold=0.80, compiled_kg=compiled_kg)

    timing['aec'] = (time.time() - aec_start) * 1000
    spinner.stop(f"Verified {len(statements)} statements against compiled knowledge", timing['aec'])
    print()

    # Per-statement verdicts
    for stmt_result in aec_result.get('statements', []):
        text = stmt_result.get('text', '')
        # Truncate smartly
        if len(text) > 50:
            text = text[:50] + "..."
        category = stmt_result.get('category', 'unknown')
        method = stmt_result.get('method', '')
        node = stmt_result.get('matched_node', stmt_result.get('matched_label', ''))
        coverage = stmt_result.get('coverage', 0)

        if category == 'grounded':
            # Determine display based on method
            if method.startswith('deterministic_'):
                # Layer 0: entity/name/number match
                match_type = method.replace('deterministic_', '')
                if '(weak)' in method:
                    # Weak name match with no concept backing - display as persona
                    symbol = f"{YELLOW}◌{RESET}"
                    print(f" {symbol}  \"{text}\"")
                    print(f"    {YELLOW}PERSONA{RESET}  {DIM}·  weak name match only{RESET}")
                elif coverage == 0 and not node:
                    # No real match - display as persona
                    symbol = f"{YELLOW}◌{RESET}"
                    print(f" {symbol}  \"{text}\"")
                    print(f"    {YELLOW}PERSONA{RESET}  {DIM}·  no verifiable content{RESET}")
                else:
                    symbol = f"{GREEN}✓{RESET}"
                    print(f" {symbol}  \"{text}\"")
                    print(f"    {DIM}Layer 0  ·  {match_type} match{RESET}")
            elif 'concept:' in method:
                # Layer 1: concept pattern match
                symbol = f"{GREEN}✓{RESET}"
                print(f" {symbol}  \"{text}\"")
                print(f"    {DIM}Layer 1  ·  {CYAN}{node}{RESET}  {DIM}·  coverage {coverage:.2f}{RESET}")
            elif 'llm' in method.lower():
                # Layer 2: LLM verification
                symbol = f"{GREEN}✓{RESET}"
                print(f" {symbol}  \"{text}\"")
                print(f"    {DIM}Layer 2  ·  {CYAN}{node}{RESET}  {DIM}·  LLM verified{RESET}")
            else:
                # Unknown method but grounded - show if valid
                if coverage > 0 and node:
                    symbol = f"{GREEN}✓{RESET}"
                    print(f" {symbol}  \"{text}\"")
                    print(f"    {DIM}Layer 1  ·  {CYAN}{node}{RESET}  {DIM}·  coverage {coverage:.2f}{RESET}")
                else:
                    # No real evidence - display as persona
                    symbol = f"{YELLOW}◌{RESET}"
                    print(f" {symbol}  \"{text}\"")
                    print(f"    {YELLOW}PERSONA{RESET}  {DIM}·  no verifiable content{RESET}")
            print()
        elif category == 'ungrounded':
            symbol = f"{RED}✗{RESET}"
            print(f" {symbol}  \"{text}\"")
            if 'antipattern' in method:
                terms = stmt_result.get('violation_terms', [])
                node = stmt_result.get('matched_node', '')
                print(f"    {RED}VIOLATION{RESET}  {DIM}·  {CYAN}{node}{RESET}  {DIM}·  terms: {terms}{RESET}")
            else:
                print(f"    {RED}UNGROUNDED{RESET}  {DIM}·  no matching node  ·  gap queued{RESET}")
            print()
        else:  # persona
            symbol = f"{YELLOW}◌{RESET}"
            print(f" {symbol}  \"{text}\"")
            print(f"    {YELLOW}PERSONA{RESET}  {DIM}·  no verifiable content{RESET}")
            print()

    # Score summary
    print(f"{'─'*WIDTH}")
    print()
    print(f"  {BOLD}SCORE{RESET}   {render_gauge(aec_result['score'], aec_result['threshold'])}")

    result_status = f"{GREEN}✓ PASS{RESET}" if aec_result['passed'] else f"{RED}✗ FAIL{RESET}"
    g = aec_result['grounded_statements']
    u = aec_result['ungrounded_statements']
    p = aec_result['persona_statements']
    print(f"  {BOLD}RESULT{RESET}  {result_status}  {DIM}·  {g} grounded  ·  {u} ungrounded  ·  {p} persona{RESET}")
    print()
    print(f"{'─'*WIDTH}")

    # ═══════════════════════════════════════════════════════════════════════════
    # ARTIFACT PRODUCED
    # ═══════════════════════════════════════════════════════════════════════════
    if not no_docx:
        print()
        print(f"{CYAN}{'━'*WIDTH}{RESET}")
        print(f"{CYAN}◈  ARTIFACT PRODUCED{RESET}")
        print(f"{CYAN}{'━'*WIDTH}{RESET}")
        print()

        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        output_path = Path("outputs") / f"docx-demo-output-{timestamp}.docx"
        success, size_kb = create_docx_file(response, output_path)

        if success:
            print_box_top()
            print_box_line(f"📄  {GREEN}{output_path}{RESET}  ({size_kb:.0f} KB)")
            print_box_empty()
            print_box_line(f"{DIM}Contents:{RESET}")
            print_box_line(f"  · Title Page")
            print_box_line(f"  · Executive Summary (Heading 1)")
            print_box_line(f"  · Milestones Table (3 columns)")
            print_box_line(f"  · Risk Register (Heading 1)")
            print_box_empty()
            print_box_line(f"{DIM}Open in Microsoft Word or Google Docs.{RESET}")
            print_box_bottom()
        else:
            print(f"  {YELLOW}Install python-docx for .docx output: pip install python-docx{RESET}")

    # ═══════════════════════════════════════════════════════════════════════════
    # TIMING SUMMARY
    # ═══════════════════════════════════════════════════════════════════════════
    total_time = (time.time() - pipeline_start) * 1000

    print()
    print(f"{BOLD}Pipeline complete in {total_time/1000:.2f}s{RESET}")
    print(f"{'─'*28}")
    print(f"  Distill    {timing.get('distill', 0):>6.0f}ms")
    print(f"  Augment    {timing.get('augment', 0):>6.0f}ms")
    if timing.get('generate', 0) > 1000:
        print(f"  Generate   {timing.get('generate', 0)/1000:>6.1f}s   {DIM}(LLM call){RESET}")
    else:
        print(f"  Generate   {timing.get('generate', 0):>6.0f}ms")
    print(f"  AEC        {timing.get('aec', 0):>6.2f}ms")
    print(f"{'─'*28}")
    print(f"{DIM}The model did not change. The skill did.{RESET}")
    print()

    return aec_result


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="AETHER Agent Demo - demonstrates DAGR pipeline with AEC verification"
    )
    parser.add_argument("--stub", action="store_true",
        help="Use hardcoded response (no API key needed)")
    parser.add_argument("--no-docx", action="store_true",
        help="Skip .docx file generation")
    parser.add_argument("--capsule", default="docx",
        help="Capsule path or short name (default: docx)")

    args = parser.parse_args()

    # Fix Windows console encoding
    if sys.platform == "win32":
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

    # Resolve capsule path
    examples_dir = Path(__file__).parent / "examples"
    capsule_path = resolve_capsule_path(args.capsule, examples_dir)

    if not capsule_path:
        print(f"{RED}Capsule not found: {args.capsule}{RESET}")
        print(f"\nAvailable capsules in examples/:")
        for d in sorted(examples_dir.iterdir()):
            if d.is_dir():
                print(f"  - {d.name}")
        sys.exit(1)

    run_demo(capsule_path, args.stub, args.no_docx)


if __name__ == "__main__":
    main()
