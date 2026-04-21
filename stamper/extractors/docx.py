"""
docx.py - Microsoft Word document extractor.
"""

from pathlib import Path
from .base import BaseExtractor


class DocxExtractor(BaseExtractor):
    """Extract content from .docx files."""

    def extract(self) -> dict:
        """
        Extract content from Word document.

        Requires: python-docx package
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return {
                "error": "dependency_missing",
                "message": "python-docx not installed. Run: pip install python-docx",
                "sections": [],
                "tables": [],
                "raw_text": ""
            }

        # Handle path or content
        path = Path(self.content) if not isinstance(self.content, str) else Path(self.content)

        try:
            doc = DocxDocument(str(path))
        except Exception as e:
            return {
                "error": "extraction_failed",
                "message": str(e),
                "sections": [],
                "tables": [],
                "raw_text": ""
            }

        sections = []
        current_section = {
            "heading": None,
            "level": 0,
            "content": [],
            "lists": [],
            "tables": []
        }

        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            if style.startswith("Heading"):
                # Save previous section
                if current_section["content"] or current_section["lists"]:
                    current_section["content"] = '\n'.join(current_section["content"])
                    sections.append(current_section)

                # Extract level from style name
                try:
                    level = int(style.replace("Heading ", "").replace("Heading", "1"))
                except ValueError:
                    level = 1

                current_section = {
                    "heading": text,
                    "level": level,
                    "content": [],
                    "lists": [],
                    "tables": []
                }
            elif style in ("List Bullet", "List Number", "List Paragraph"):
                current_section["lists"].append(text)
            else:
                current_section["content"].append(text)

        # Save final section
        if current_section["content"] or current_section["lists"]:
            current_section["content"] = '\n'.join(current_section["content"])
            sections.append(current_section)

        # Extract tables
        tables = []
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            header_row = rows[0] if rows else []
            has_header = any(
                kw in cell.upper()
                for cell in header_row
                for kw in ["MUST", "SHALL", "RULE", "CONSTRAINT", "REQUIREMENT"]
            )
            tables.append({
                "rows": rows,
                "has_header": has_header,
                "header_row": header_row
            })

        # Build raw text
        raw_text = '\n'.join(para.text for para in doc.paragraphs if para.text.strip())

        return {
            "frontmatter": {},
            "sections": sections,
            "prohibitions_in_frontmatter": [],
            "tables": tables,
            "raw_text": raw_text
        }
